from __future__ import annotations

import json
import os
import re
from pathlib import Path

import pytest
from alembic import command
from alembic.config import Config
from docx import Document
from sqlalchemy import create_engine, text
from sqlalchemy.engine import URL, make_url
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from job_agent_runtime.orchestration.manager import ManagerPipelineRunner
from job_agent_runtime.agents.stub_agents import DefaultStubAgentFactory
from job_agent_runtime.agents.base_agent import BaseAgent
from job_platform.config import clear_settings_cache, get_settings
from job_agent_runtime.worker.watcher import WatcherService


ROOT_DIR = Path(__file__).resolve().parents[3]
DEFAULT_TEST_DATABASE_URL = "postgresql+asyncpg://postgres:password@localhost:5432/jobagent_test"


def _async_to_sync_url(database_url: str) -> str:
    if database_url.startswith("postgresql+asyncpg://"):
        return database_url.replace("postgresql+asyncpg://", "postgresql+psycopg2://", 1)
    if database_url.startswith("postgres+asyncpg://"):
        return database_url.replace("postgres+asyncpg://", "postgres+psycopg2://", 1)
    return database_url


def _ensure_database_exists(sync_test_url: str) -> None:
    test_url = make_url(sync_test_url)
    database_name = test_url.database or ""
    if not re.fullmatch(r"[A-Za-z0-9_]+", database_name):
        raise RuntimeError(f"Unsafe test database name: {database_name!r}")

    admin_url: URL = test_url.set(database="postgres")
    admin_engine = create_engine(admin_url, isolation_level="AUTOCOMMIT")
    try:
        with admin_engine.connect() as conn:
            conn.execute(text("SELECT 1"))
            exists = conn.execute(
                text("SELECT 1 FROM pg_database WHERE datname = :db_name"),
                {"db_name": database_name},
            ).scalar_one_or_none()
            if not exists:
                conn.execute(text(f'CREATE DATABASE "{database_name}"'))
    except Exception as exc:  # pragma: no cover
        pytest.skip(f"PostgreSQL not available for integration tests: {exc}")
    finally:
        admin_engine.dispose()


def _reset_public_schema(sync_test_url: str) -> None:
    engine = create_engine(sync_test_url, isolation_level="AUTOCOMMIT")
    try:
        with engine.connect() as conn:
            conn.execute(text("DROP SCHEMA IF EXISTS public CASCADE"))
            conn.execute(text("CREATE SCHEMA public"))
            conn.execute(text("GRANT ALL ON SCHEMA public TO postgres"))
            conn.execute(text("GRANT ALL ON SCHEMA public TO public"))
    finally:
        engine.dispose()


def _run_alembic(command_name: str) -> None:
    config = Config(str(ROOT_DIR / "alembic.ini"))
    config.set_main_option("script_location", str(ROOT_DIR / "alembic"))
    if command_name == "upgrade":
        command.upgrade(config, "head")
        return
    if command_name == "downgrade":
        command.downgrade(config, "base")
        return
    raise ValueError(f"Unsupported Alembic command: {command_name}")


def _set_required_runtime_env(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    database_url: str,
) -> None:
    data_dir = tmp_path / "data"
    output_dir = tmp_path / "output"
    resume_library_dir = data_dir / "resume-library"
    resume_tracks_dir = data_dir / "resume-tracks"
    resume_docx_tracks_dir = data_dir / "resume-docx-tracks"
    skills_dir = ROOT_DIR / "apps" / "agent-runtime" / "skills"
    data_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    resume_library_dir.mkdir(parents=True, exist_ok=True)
    resume_tracks_dir.mkdir(parents=True, exist_ok=True)
    resume_docx_tracks_dir.mkdir(parents=True, exist_ok=True)

    base_resume_docx = data_dir / "base_resume.docx"
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Old summary text.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, SQL")
    doc.add_heading("Relevant Experience", level=1)
    doc.add_paragraph("Built ML services on AWS.")
    doc.save(str(base_resume_docx))

    base_resume_text = data_dir / "base_resume.md"
    credentials = data_dir / "credentials.json"
    base_resume_text.write_text("Candidate resume text", encoding="utf-8")
    credentials.write_text("{}", encoding="utf-8")
    track_payload = {
        "display_name": "Python ML Track",
        "raw_text": "Summary\nPython ML engineer\nSkills\nPython AWS LLM\nExperience\nRecent ML role",
        "normalized_text": "Summary\nPython ML engineer\nSkills\nPython AWS LLM\nExperience\nRecent ML role",
        "sections": {
            "summary": "Python ML engineer",
            "skills": "Python\nAWS\nLLM",
            "experience_recent_role": "Built ML services on AWS.",
            "education": "MS Computer Science",
        },
        "role_bias": ["ai_ml", "backend_python", "cloud_platform"],
        "keywords": ["python", "aws", "llm", "machine learning"],
    }
    for suffix in ("python_ml", "data_platform", "general_backend"):
        payload = dict(track_payload)
        payload["track_id"] = f"resume_track_{suffix}"
        payload["source_pdf_path"] = str(resume_library_dir / f"resume_track_{suffix}.pdf")
        (resume_tracks_dir / f"resume_track_{suffix}.json").write_text(json.dumps(payload), encoding="utf-8")
        doc.save(str(resume_docx_tracks_dir / f"resume_track_{suffix}.docx"))

    env = {
        "ANTHROPIC_API_KEY": "sk-ant-test",
        "MANAGER_MODEL": "claude-opus-4-6",
        "RESEARCH_MODEL": "claude-sonnet-4-6",
        "RESUME_EDITOR_MODEL": "claude-sonnet-4-6",
        "GMAIL_AGENT_MODEL": "claude-sonnet-4-6",
        "WHATSAPP_MSG_MODEL": "claude-sonnet-4-6",
        "DATABASE_URL": database_url,
        "TEST_DATABASE_URL": database_url,
        "WAHA_BASE_URL": "http://localhost:3000",
        "WAHA_SESSION": "default",
        "WAHA_API_KEY": "waha-test",
        "WHATSAPP_GROUP_IDS": "GROUP1@g.us",
        "POLL_INTERVAL_SECONDS": "1800",
        "GMAIL_CREDENTIALS_PATH": str(credentials),
        "GMAIL_TOKEN_PATH": str(data_dir / "token.json"),
        "SENDER_EMAIL": "sender@example.com",
        "SENDER_NAME": "Pranay",
        "MIN_RELEVANCE_SCORE": "6",
        "MIN_ATS_SCORE": "65",
        "MAX_RESUME_EDIT_ITERATIONS": "2",
        "BASE_RESUME_DOCX": str(base_resume_docx),
        "BASE_RESUME_TEXT": str(base_resume_text),
        "RESUME_LIBRARY_DIR": str(resume_library_dir),
        "RESUME_TRACKS_DIR": str(resume_tracks_dir),
        "RESUME_DOCX_TRACKS_DIR": str(resume_docx_tracks_dir),
        "OUTPUT_DIR": str(output_dir),
        "SKILLS_DIR": str(skills_dir),
        "LOG_LEVEL": "INFO",
        "LOG_FORMAT": "json",
        "OTEL_ENDPOINT": "http://localhost:4317",
    }
    for key, value in env.items():
        monkeypatch.setenv(key, value)


@pytest.mark.asyncio
async def test_real_content_pipeline_with_docx_attachment_delivery(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    async_test_url = os.getenv("TEST_DATABASE_URL", DEFAULT_TEST_DATABASE_URL)
    sync_test_url = _async_to_sync_url(async_test_url)

    _ensure_database_exists(sync_test_url)
    _reset_public_schema(sync_test_url)
    _set_required_runtime_env(monkeypatch, tmp_path, async_test_url)

    clear_settings_cache()
    _run_alembic("upgrade")
    upgraded = True

    async_engine = create_async_engine(async_test_url)
    session_factory = async_sessionmaker(async_engine, class_=AsyncSession, expire_on_commit=False)

    async def _fake_call_model(self, messages, trace_id, tools=None, max_tokens=2048):  # noqa: ANN001, ARG002
        content = str(messages[-1].get("content", "")).lower()
        if "evaluate whether this whatsapp message" in content:
            payload = {
                "relevant": True,
                "score": 8,
                "job_title": "ML Engineer",
                "company": "Acme",
                "job_summary": "Python and ML role",
                "poster_email": "recruiter@example.com",
                "poster_number": "+15550001111",
                "discard_reason": None,
                "relevance_reason": "Strong relevance",
            }
            return {"text": json.dumps(payload)}

        if "shortlisted_tracks" in content and "selected_resume_track" in content:
            payload = {
                "add_items": [
                    {
                        "section": "summary",
                        "action": "Add LLM impact line",
                        "reason": "Role asks for LLM experience",
                        "priority": 1,
                    }
                ],
                "remove_items": [],
                "keywords_to_inject": ["Python", "LLM"],
                "sections_to_edit": ["summary", "skills", "experience_recent_role"],
                "ats_score_estimate_before": 55,
                "ats_score_estimate_after": 78,
                "research_reasoning": "Focus on summary alignment.",
                "selected_resume_track": "resume_track_python_ml",
                "selected_resume_source_pdf": "data/resume-library/resume_track_python_ml.pdf",
                "selected_resume_match_reason": "Strongest Python and ML evidence density.",
                "experience_target_section": "experience_recent_role",
                "summary_focus": "Reframe the summary around Python and LLM delivery impact.",
                "skills_gap_notes": ["Surface grounded LLM and FastAPI terminology."],
                "hard_gaps": [],
                "edit_scope": ["summary", "skills", "experience_recent_role"],
            }
            return {"text": json.dumps(payload), "input_tokens": 30, "output_tokens": 40, "latency_ms": 5}

        if "apply targeted edits to resume sections only" in content:
            payload = {
                "edited_sections": {
                    "summary": "ML Engineer with Python and LLM delivery impact."
                },
                "changes_applied": ["Updated summary for target role"],
                "evaluation": {
                    "ats_score_before": 55,
                    "ats_score_after": 80,
                    "checklist_passed": True,
                    "iterations": 1,
                },
            }
            return {"text": json.dumps(payload), "input_tokens": 35, "output_tokens": 55, "latency_ms": 8}

        if "quality-gate this resume iteration" in content:
            payload = {"pass": True, "reason": "Looks good", "feedback": ""}
            return {"text": json.dumps(payload)}

        return {"text": "{}"}

    monkeypatch.setattr(BaseAgent, "_call_model", _fake_call_model)

    try:
        sync_engine = create_engine(sync_test_url)
        try:
            with sync_engine.begin() as conn:
                conn.execute(
                    text(
                        """
                        INSERT INTO whatsapp_messages (group_id, sender_number, message_text, message_hash)
                        VALUES ('GROUP1@g.us', '+15550000001', 'Python ML role recruiter@example.com', 'wm_step8_hash_1')
                        """
                    )
                )
        finally:
            sync_engine.dispose()

        settings = get_settings()
        pipeline_runner = ManagerPipelineRunner(
            session_factory=session_factory,
            settings=settings,
            agent_factory=DefaultStubAgentFactory(settings=settings),
        )
        watcher = WatcherService(settings=settings, session_factory=session_factory)
        summary = await watcher.run_tick(pipeline_runner=pipeline_runner)
        assert summary["processed_count"] == 1
        assert summary["error_count"] == 0

        sync_engine = create_engine(sync_test_url)
        try:
            with sync_engine.connect() as conn:
                row = conn.execute(
                    text(
                        """
                        SELECT pr.status, rv.docx_path, rv.attachment_path
                        FROM pipeline_runs pr
                        JOIN resume_versions rv ON rv.trace_id = pr.trace_id
                        ORDER BY pr.created_at DESC, rv.version_number DESC
                        LIMIT 1
                        """
                    )
                ).first()
                assert row is not None
                assert row[0] == "sent"
                assert row[1]
                assert row[2]
                assert Path(str(row[1])).is_file()
                assert Path(str(row[2])).is_file()
                assert Path(str(row[1])) == Path(str(row[2]))
        finally:
            sync_engine.dispose()
    finally:
        await async_engine.dispose()
        clear_settings_cache()
        if upgraded:
            _run_alembic("downgrade")



