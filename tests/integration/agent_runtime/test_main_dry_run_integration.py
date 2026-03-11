from __future__ import annotations

import json
import re
from pathlib import Path

import pytest
from alembic import command
from alembic.config import Config
from docx import Document
from sqlalchemy import create_engine, text
from sqlalchemy.engine import URL, make_url
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

import job_agent_runtime.main as main_module
from job_agent_runtime.orchestration import manager as manager_module
from job_platform.config import clear_settings_cache, get_settings


ROOT_DIR = Path(__file__).resolve().parents[3]
DEFAULT_TEST_DATABASE_URL = "postgresql+asyncpg://postgres:password@localhost:5432/jobagent_step910_dryrun_test"


def _async_to_sync_url(database_url: str) -> str:
    if database_url.startswith("postgresql+asyncpg://"):
        return database_url.replace("postgresql+asyncpg://", "postgresql+psycopg2://", 1)
    if database_url.startswith("postgres+asyncpg://"):
        return database_url.replace("postgres+asyncpg://", "postgres+psycopg2://", 1)
    return database_url


def _ensure_database_exists(sync_test_url: str) -> None:
    test_url = make_url(sync_test_url)
    database_name = test_url.database or ""
    if not re.fullmatch(r"[A-Za-z0-9_]+", database_name):
        raise RuntimeError(f"Unsafe test database name: {database_name!r}")

    admin_url: URL = test_url.set(database="postgres")
    admin_engine = create_engine(admin_url, isolation_level="AUTOCOMMIT")
    try:
        with admin_engine.connect() as conn:
            conn.execute(text("SELECT 1"))
            exists = conn.execute(
                text("SELECT 1 FROM pg_database WHERE datname = :db_name"),
                {"db_name": database_name},
            ).scalar_one_or_none()
            if not exists:
                conn.execute(text(f'CREATE DATABASE "{database_name}"'))
    except Exception as exc:  # pragma: no cover
        pytest.skip(f"PostgreSQL not available for integration tests: {exc}")
    finally:
        admin_engine.dispose()


def _reset_public_schema(sync_test_url: str) -> None:
    db_engine = create_engine(sync_test_url, isolation_level="AUTOCOMMIT")
    try:
        with db_engine.connect() as conn:
            conn.execute(text("DROP SCHEMA IF EXISTS public CASCADE"))
            conn.execute(text("CREATE SCHEMA IF NOT EXISTS public"))
            conn.execute(text("GRANT ALL ON SCHEMA public TO postgres"))
            conn.execute(text("GRANT ALL ON SCHEMA public TO public"))
    finally:
        db_engine.dispose()


def _run_alembic(command_name: str) -> None:
    config = Config(str(ROOT_DIR / "alembic.ini"))
    config.set_main_option("script_location", str(ROOT_DIR / "alembic"))
    if command_name == "upgrade":
        command.upgrade(config, "head")
        return
    if command_name == "downgrade":
        command.downgrade(config, "base")
        return
    raise ValueError(f"Unsupported Alembic command: {command_name}")


def _set_required_runtime_env(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    database_url: str,
) -> None:
    data_dir = tmp_path / "data"
    output_dir = tmp_path / "output"
    resume_library_dir = data_dir / "resume-library"
    resume_tracks_dir = data_dir / "resume-tracks"
    resume_docx_tracks_dir = data_dir / "resume-docx-tracks"
    skills_dir = ROOT_DIR / "apps" / "agent-runtime" / "skills"
    data_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    resume_library_dir.mkdir(parents=True, exist_ok=True)
    resume_tracks_dir.mkdir(parents=True, exist_ok=True)
    resume_docx_tracks_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "resumes").mkdir(parents=True, exist_ok=True)
    base_docx = data_dir / "base_resume.docx"
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Placeholder summary")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, SQL")
    doc.add_heading("Relevant Experience", level=1)
    doc.add_paragraph("Built backend and AI services.")
    doc.save(str(base_docx))
    (data_dir / "base_resume.md").write_text("placeholder", encoding="utf-8")
    (data_dir / "credentials.json").write_text("{}", encoding="utf-8")
    track_payload = {
        "display_name": "Python ML Track",
        "raw_text": "Summary\nPython ML engineer\nSkills\nPython AWS LLM\nExperience\nRecent ML role",
        "normalized_text": "Summary\nPython ML engineer\nSkills\nPython AWS LLM\nExperience\nRecent ML role",
        "sections": {
            "summary": "Python ML engineer",
            "skills": "Python\nAWS\nLLM",
            "experience_recent_role": "Built ML services on AWS.",
            "education": "MS Computer Science",
        },
        "role_bias": ["ai_ml", "backend_python", "cloud_platform"],
        "keywords": ["python", "aws", "llm", "machine learning"],
    }
    for suffix in ("python_ml", "data_platform", "general_backend"):
        payload = dict(track_payload)
        payload["track_id"] = f"resume_track_{suffix}"
        payload["source_pdf_path"] = str(resume_library_dir / f"resume_track_{suffix}.pdf")
        (resume_tracks_dir / f"resume_track_{suffix}.json").write_text(json.dumps(payload), encoding="utf-8")
        doc.save(str(resume_docx_tracks_dir / f"resume_track_{suffix}.docx"))

    env = {
        "ANTHROPIC_API_KEY": "sk-ant-test",
        "MANAGER_MODEL": "claude-opus-4-6",
        "RESEARCH_MODEL": "claude-sonnet-4-6",
        "RESUME_EDITOR_MODEL": "claude-sonnet-4-6",
        "GMAIL_AGENT_MODEL": "claude-sonnet-4-6",
        "WHATSAPP_MSG_MODEL": "claude-sonnet-4-6",
        "DATABASE_URL": database_url,
        "TEST_DATABASE_URL": database_url,
        "WAHA_BASE_URL": "http://localhost:3000",
        "WAHA_SESSION": "default",
        "WAHA_API_KEY": "waha-test",
        "WHATSAPP_GROUP_IDS": "GROUP1@g.us",
        "POLL_INTERVAL_SECONDS": "1800",
        "GMAIL_CREDENTIALS_PATH": str(data_dir / "credentials.json"),
        "GMAIL_TOKEN_PATH": str(data_dir / "token.json"),
        "SENDER_EMAIL": "sender@example.com",
        "SENDER_NAME": "Pranay",
        "MIN_RELEVANCE_SCORE": "6",
        "MIN_ATS_SCORE": "65",
        "MAX_RESUME_EDIT_ITERATIONS": "2",
        "BASE_RESUME_DOCX": str(base_docx),
        "BASE_RESUME_TEXT": str(data_dir / "base_resume.md"),
        "RESUME_LIBRARY_DIR": str(resume_library_dir),
        "RESUME_TRACKS_DIR": str(resume_tracks_dir),
        "RESUME_DOCX_TRACKS_DIR": str(resume_docx_tracks_dir),
        "OUTPUT_DIR": str(output_dir),
        "SKILLS_DIR": str(skills_dir),
        "LOG_LEVEL": "INFO",
        "LOG_FORMAT": "json",
        "OTEL_ENDPOINT": "http://localhost:4317",
    }
    for key, value in env.items():
        monkeypatch.setenv(key, value)


@pytest.mark.asyncio
async def test_dry_run_rolls_back_transient_rows(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    async_test_url = DEFAULT_TEST_DATABASE_URL
    sync_test_url = _async_to_sync_url(async_test_url)

    _ensure_database_exists(sync_test_url)
    _reset_public_schema(sync_test_url)
    _set_required_runtime_env(monkeypatch, tmp_path, async_test_url)

    clear_settings_cache()
    _run_alembic("upgrade")
    upgraded = True

    settings = get_settings()
    async_engine = create_async_engine(async_test_url)
    session_factory = async_sessionmaker(async_engine, class_=AsyncSession, expire_on_commit=False)
    monkeypatch.setattr(main_module, "AsyncSessionLocal", session_factory)

    async def _fake_fetch(_settings):  # noqa: ANN001
        return [
            {
                "group_id": "GROUP1@g.us",
                "sender_number": "+15550001111",
                "text": "Dry run sample",
                "timestamp": 1710000000,
            }
        ]

    async def _fake_manager_run(self, message, trace_id):  # noqa: ANN001, ARG002
        _ = message
        return {"trace_id": str(trace_id), "action": "dry_run_ready"}

    monkeypatch.setattr(main_module, "_fetch_dry_run_samples", _fake_fetch)
    monkeypatch.setattr(manager_module.ManagerAgent, "run", _fake_manager_run)

    try:
        exit_code = await main_module._run_dry_run(settings)
        assert exit_code == 0

        db_engine = create_engine(sync_test_url)
        try:
            with db_engine.connect() as conn:
                message_count = conn.execute(text("SELECT count(*) FROM whatsapp_messages")).scalar_one()
                pipeline_count = conn.execute(text("SELECT count(*) FROM pipeline_runs")).scalar_one()
                resume_count = conn.execute(text("SELECT count(*) FROM resume_versions")).scalar_one()
                outbox_count = conn.execute(text("SELECT count(*) FROM outbox")).scalar_one()
                trace_count = conn.execute(text("SELECT count(*) FROM agent_traces")).scalar_one()
                assert message_count == 0
                assert pipeline_count == 0
                assert resume_count == 0
                assert outbox_count == 0
                assert trace_count == 0
        finally:
            db_engine.dispose()
    finally:
        await async_engine.dispose()
        clear_settings_cache()
        if upgraded:
            _run_alembic("downgrade")



