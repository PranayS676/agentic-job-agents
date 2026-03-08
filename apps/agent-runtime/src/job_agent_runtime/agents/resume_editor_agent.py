from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.text.paragraph import Paragraph
from sqlalchemy.ext.asyncio import AsyncSession

from job_agent_runtime.agents.base_agent import BaseAgent
from job_platform.config import Settings, get_settings
from job_platform.tracer import AgentTracer

from .contracts import ResearchOutput, ResumeEditOutput


class ResumeEditorAgent(BaseAgent):
    def __init__(
        self,
        *,
        db_session: AsyncSession,
        tracer: AgentTracer,
        settings: Settings | None = None,
    ) -> None:
        self.settings = settings or get_settings()
        super().__init__(
            skill_path="skills/resume-editor",
            model=self.settings.resume_editor_model,
            db_session=db_session,
            tracer=tracer,
        )

    async def run(
        self,
        research_output: ResearchOutput,
        trace_id: UUID,
        job_context: dict,
        version_number: int,
        feedback: str | None = None,
    ) -> ResumeEditOutput:
        if version_number < 1:
            raise ValueError("version_number must be >= 1")

        if self.settings.base_resume_docx is None:
            raise FileNotFoundError("BASE_RESUME_DOCX is not configured")
        base_docx_path = self.settings.resolve_path(self.settings.base_resume_docx)
        if not base_docx_path.is_file():
            raise FileNotFoundError(f"BASE_RESUME_DOCX file not found: {base_docx_path}")

        document = Document(str(base_docx_path))
        section_texts, section_paragraphs = self._extract_sections(document)
        sections_to_edit = [self._canonical_key(section) for section in research_output["sections_to_edit"]]
        target_sections = {
            section_key: section_texts.get(section_key, "")
            for section_key in sections_to_edit
        }

        prompt_payload = {
            "research_output": research_output,
            "target_sections": target_sections,
            "feedback": feedback,
        }
        prompt = (
            "Apply targeted edits to resume sections only. Return strict JSON with keys "
            "'edited_sections', 'changes_applied', and 'evaluation'.\n"
            f"Input:\n{json.dumps(prompt_payload, ensure_ascii=True, indent=2)}"
        )
        model_result = await self._call_model(
            messages=[{"role": "user", "content": prompt}],
            trace_id=trace_id,
            max_tokens=4096,
        )
        parsed = self._parse_json(model_result["text"])
        edited_sections, changes_applied, evaluation = self._coerce_editor_output(parsed)

        applied_sections, skipped_sections = self._apply_section_edits(
            document=document,
            section_paragraphs=section_paragraphs,
            sections_to_edit=sections_to_edit,
            edited_sections=edited_sections,
        )

        if self.settings.output_dir is None:
            raise RuntimeError("OUTPUT_DIR is not configured")
        output_dir = self.settings.resolve_path(self.settings.output_dir) / "resumes"
        output_dir.mkdir(parents=True, exist_ok=True)
        company_slug = self._slugify(str(job_context.get("company") or "unknown"))
        job_slug = self._slugify(str(job_context.get("job_title") or "unknown"))
        output_docx = output_dir / f"{company_slug}_{job_slug}_{str(trace_id)[:8]}_v{version_number}.docx"
        document.save(str(output_docx))

        before_score = self._coerce_optional_int(
            evaluation.get("ats_score_before"),
            fallback=research_output["ats_score_estimate_before"],
        )
        after_score = self._coerce_optional_int(evaluation.get("ats_score_after"), fallback=None)
        if after_score is None:
            resume_text = self._document_to_text(document)
            scorer_score = self._score_with_script(
                resume_text=resume_text,
                keywords=research_output["keywords_to_inject"],
            )
            after_score = scorer_score if scorer_score is not None else max(before_score or 0, 0)

        checklist_passed = bool(
            evaluation.get("checklist_passed", after_score >= (before_score or 0))
        )
        evaluation_summary = str(
            evaluation.get("summary")
            or evaluation.get("reason")
            or f"checklist_passed={checklist_passed}; ats={before_score}->{after_score}"
        )

        return {
            "docx_path": str(output_docx),
            "changes_made": {
                "edited_sections": edited_sections,
                "applied_sections": applied_sections,
                "skipped_sections": skipped_sections,
                "changes_applied": changes_applied,
                "evaluation": evaluation,
            },
            "ats_score_before": before_score or 0,
            "ats_score_after": after_score,
            "evaluator_passed": checklist_passed,
            "evaluation_summary": evaluation_summary,
        }

    def _extract_sections(
        self,
        document: Document,
    ) -> tuple[dict[str, str], dict[str, list[int]]]:
        section_paragraphs: dict[str, list[int]] = {}
        current_section = "body"
        section_paragraphs[current_section] = []
        seen_sections: set[str] = {current_section}

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            if self._is_section_heading(paragraph):
                base_key = self._canonical_key(text)
                section_key = base_key or "section"
                suffix = 2
                while section_key in seen_sections:
                    section_key = f"{base_key}_{suffix}"
                    suffix += 1
                seen_sections.add(section_key)
                current_section = section_key
                section_paragraphs.setdefault(current_section, [])
                continue

            section_paragraphs.setdefault(current_section, []).append(index)

        if all(len(indexes) == 0 for indexes in section_paragraphs.values()):
            all_indexes = [idx for idx, p in enumerate(document.paragraphs) if p.text.strip()]
            section_paragraphs = {"body": all_indexes}

        section_texts: dict[str, str] = {}
        for section_key, paragraph_indexes in section_paragraphs.items():
            lines = [document.paragraphs[idx].text.strip() for idx in paragraph_indexes if document.paragraphs[idx].text.strip()]
            section_texts[section_key] = "\n".join(lines)
        return section_texts, section_paragraphs

    def _is_section_heading(self, paragraph: Paragraph) -> bool:
        text = paragraph.text.strip()
        if not text:
            return False
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.startswith("Heading"):
            return True
        runs = [run for run in paragraph.runs if run.text.strip()]
        if runs and all(run.bold for run in runs) and len(text) <= 80:
            return True
        return False

    def _coerce_editor_output(
        self,
        payload: dict[str, Any],
    ) -> tuple[dict[str, str], list[str], dict[str, Any]]:
        required = {"edited_sections", "changes_applied", "evaluation"}
        missing = sorted(required - set(payload))
        if missing:
            raise ValueError(f"Resume editor output missing required keys: {', '.join(missing)}")

        edited_sections_raw = payload.get("edited_sections")
        if not isinstance(edited_sections_raw, dict):
            raise ValueError("edited_sections must be an object")
        edited_sections: dict[str, str] = {}
        for key, value in edited_sections_raw.items():
            key_text = str(key).strip()
            value_text = str(value or "").strip()
            if not key_text or not value_text:
                continue
            edited_sections[key_text] = value_text

        changes_applied_raw = payload.get("changes_applied")
        if not isinstance(changes_applied_raw, list):
            raise ValueError("changes_applied must be a list")
        changes_applied = [str(item).strip() for item in changes_applied_raw if str(item).strip()]

        evaluation_raw = payload.get("evaluation")
        if not isinstance(evaluation_raw, dict):
            raise ValueError("evaluation must be an object")

        return edited_sections, changes_applied, evaluation_raw

    def _apply_section_edits(
        self,
        *,
        document: Document,
        section_paragraphs: dict[str, list[int]],
        sections_to_edit: list[str],
        edited_sections: dict[str, str],
    ) -> tuple[dict[str, str], dict[str, str]]:
        target_set = {self._canonical_key(section) for section in sections_to_edit}
        key_map = {self._canonical_key(key): key for key in section_paragraphs}

        applied: dict[str, str] = {}
        skipped: dict[str, str] = {}

        for raw_section, edited_text in edited_sections.items():
            section_key = self._canonical_key(raw_section)
            if section_key not in target_set:
                skipped[raw_section] = "section not in sections_to_edit"
                continue

            resolved_key = key_map.get(section_key)
            if resolved_key is None:
                skipped[raw_section] = "section not found in base resume"
                continue

            paragraph_indexes = section_paragraphs.get(resolved_key, [])
            if not paragraph_indexes:
                skipped[raw_section] = "no paragraph content for section"
                continue

            first_idx = paragraph_indexes[0]
            document.paragraphs[first_idx].text = edited_text
            for idx in paragraph_indexes[1:]:
                document.paragraphs[idx].text = ""
            applied[resolved_key] = edited_text

        return applied, skipped

    def _document_to_text(self, document: Document) -> str:
        return "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())

    def _score_with_script(self, *, resume_text: str, keywords: list[str]) -> int | None:
        if self.settings.skills_dir is None:
            return None
        script_path = self.settings.resolve_path(self.settings.skills_dir) / "resume-editor" / "scripts" / "ats_scorer.py"
        if not script_path.is_file():
            return None

        with tempfile.NamedTemporaryFile("w", encoding="utf-8", delete=False, suffix=".txt") as tmp_file:
            tmp_path = Path(tmp_file.name)
            tmp_file.write(resume_text)

        try:
            command = [
                sys.executable,
                str(script_path),
                "--resume-file",
                str(tmp_path),
                "--keywords",
                ",".join(keywords),
            ]
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=20,
            )
            if result.returncode != 0:
                return None
            payload = json.loads(result.stdout.strip() or "{}")
            score = payload.get("score")
            return self._coerce_optional_int(score, fallback=None)
        except Exception:
            return None
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass

    def _coerce_optional_int(self, value: Any, fallback: int | None) -> int | None:
        if value is None:
            return fallback
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return fallback
        return max(0, min(100, parsed))

    def _canonical_key(self, value: str) -> str:
        cleaned = "".join(ch.lower() if ch.isalnum() else "_" for ch in value)
        while "__" in cleaned:
            cleaned = cleaned.replace("__", "_")
        return cleaned.strip("_")

    def _slugify(self, value: str) -> str:
        return self._canonical_key(value) or "unknown"

