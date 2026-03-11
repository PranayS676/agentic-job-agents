from __future__ import annotations

import json
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from docx import Document

from job_agent_runtime.agents.resume_editor_agent import ResumeEditorAgent
from job_platform.config import clear_settings_cache


def _set_required_env(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    data_dir = tmp_path / "data"
    output_dir = tmp_path / "output"
    skills_dir = tmp_path / "skills"
    resume_docx_tracks_dir = data_dir / "resume-docx-tracks"
    data_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    skills_dir.mkdir(parents=True, exist_ok=True)
    resume_docx_tracks_dir.mkdir(parents=True, exist_ok=True)

    base_docx = data_dir / "base_resume.docx"
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Old summary text.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, SQL")
    doc.add_heading("Relevant Experience", level=1)
    doc.add_paragraph("Built production ML services.")
    doc.save(str(base_docx))
    track_docx = resume_docx_tracks_dir / "resume_track_python_ml.docx"
    doc.save(str(track_docx))

    (data_dir / "base_resume.md").write_text("resume text", encoding="utf-8")
    (data_dir / "credentials.json").write_text("{}", encoding="utf-8")

    env = {
        "ANTHROPIC_API_KEY": "sk-ant-test",
        "MANAGER_MODEL": "claude-opus-4-6",
        "RESEARCH_MODEL": "claude-sonnet-4-6",
        "RESUME_EDITOR_MODEL": "claude-sonnet-4-6",
        "GMAIL_AGENT_MODEL": "claude-sonnet-4-6",
        "WHATSAPP_MSG_MODEL": "claude-sonnet-4-6",
        "DATABASE_URL": "postgresql+asyncpg://postgres:password@localhost:5432/jobagent",
        "WAHA_BASE_URL": "http://localhost:3000",
        "WAHA_SESSION": "default",
        "WAHA_API_KEY": "waha-test",
        "WHATSAPP_GROUP_IDS": "GROUP1@g.us",
        "POLL_INTERVAL_SECONDS": "30",
        "GMAIL_CREDENTIALS_PATH": str(data_dir / "credentials.json"),
        "GMAIL_TOKEN_PATH": str(data_dir / "token.json"),
        "SENDER_EMAIL": "sender@example.com",
        "SENDER_NAME": "Pranay",
        "MIN_RELEVANCE_SCORE": "6",
        "MIN_ATS_SCORE": "65",
        "MAX_RESUME_EDIT_ITERATIONS": "2",
        "BASE_RESUME_DOCX": str(base_docx),
        "BASE_RESUME_TEXT": str(data_dir / "base_resume.md"),
        "RESUME_DOCX_TRACKS_DIR": str(resume_docx_tracks_dir),
        "OUTPUT_DIR": str(output_dir),
        "SKILLS_DIR": str(skills_dir),
        "LOG_LEVEL": "INFO",
        "LOG_FORMAT": "json",
        "OTEL_ENDPOINT": "http://localhost:4317",
    }
    for key, value in env.items():
        monkeypatch.setenv(key, value)


def _create_skill_files(tmp_path: Path) -> None:
    skill_dir = tmp_path / "skills" / "resume-editor"
    references_dir = skill_dir / "references"
    scripts_dir = skill_dir / "scripts"
    references_dir.mkdir(parents=True, exist_ok=True)
    scripts_dir.mkdir(parents=True, exist_ok=True)
    (skill_dir / "SKILL.md").write_text("# Resume Editor Agent", encoding="utf-8")
    (references_dir / "style_rules.md").write_text("Keep style concise.", encoding="utf-8")
    (references_dir / "section_rules.md").write_text("Edit only summary, skills, and one experience section.", encoding="utf-8")
    (references_dir / "before_after_examples.md").write_text("Before and after examples.", encoding="utf-8")
    (references_dir / "forbidden_edits.md").write_text("Do not fabricate cloud experience.", encoding="utf-8")
    (scripts_dir / "ats_scorer.py").write_text(
        "import json\nimport argparse\n"
        "p=argparse.ArgumentParser();p.add_argument('--resume-file');p.add_argument('--keywords');a=p.parse_args();"
        "print(json.dumps({'score':77,'matched_keywords':['python'],'total_keywords':1}))",
        encoding="utf-8",
    )


def _build_agent(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> ResumeEditorAgent:
    _set_required_env(monkeypatch, tmp_path)
    _create_skill_files(tmp_path)
    monkeypatch.chdir(tmp_path)

    import job_agent_runtime.agents.base_agent as base_agent_module

    monkeypatch.setattr(
        base_agent_module,
        "AsyncAnthropic",
        lambda *args, **kwargs: SimpleNamespace(messages=SimpleNamespace(create=AsyncMock())),  # noqa: ARG005
    )
    db_session = SimpleNamespace(execute=AsyncMock(), flush=AsyncMock())
    tracer = SimpleNamespace(trace=AsyncMock())
    return ResumeEditorAgent(db_session=db_session, tracer=tracer)


def _research_output() -> dict:
    return {
        "add_items": [
            {
                "section": "summary",
                "action": "Add ML impact line",
                "reason": "Role asks for ML",
                "priority": 1,
            }
        ],
        "remove_items": [],
        "keywords_to_inject": ["Python", "LLM"],
        "sections_to_edit": ["summary", "skills", "experience_recent_role"],
        "ats_score_estimate_before": 50,
        "ats_score_estimate_after": 72,
        "research_reasoning": "Focus on ML impact in summary.",
        "selected_resume_track": "resume_track_python_ml",
        "selected_resume_source_pdf": "data/resume-library/resume_track_python_ml.pdf",
        "selected_resume_match_reason": "Strongest Python and ML evidence density.",
        "experience_target_section": "experience_recent_role",
        "summary_focus": "Align the summary around Python and ML delivery impact.",
        "skills_gap_notes": ["Surface grounded FastAPI and LLM keywords."],
        "hard_gaps": [],
        "edit_scope": ["summary", "skills", "experience_recent_role"],
    }


@pytest.fixture(autouse=True)
def _clear_cache():
    clear_settings_cache()
    yield
    clear_settings_cache()


@pytest.mark.asyncio
async def test_run_applies_only_target_sections_and_uses_filename_contract(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    agent = _build_agent(monkeypatch, tmp_path)
    trace_id = uuid4()
    model_payload = {
        "edited_sections": {
            "summary": "Updated summary with Python and LLM impact.",
            "education": "SHOULD_NOT_APPLY",
        },
        "changes_applied": ["Updated summary line"],
        "evaluation": {
            "ats_score_before": 50,
            "ats_score_after": 74,
            "checklist_passed": True,
        },
    }
    agent._call_model = AsyncMock(return_value={"text": json.dumps(model_payload)})  # type: ignore[method-assign]

    result = await agent.run(
        research_output=_research_output(),
        trace_id=trace_id,
        job_context={"company": "Acme Inc", "job_title": "ML Engineer", "relevance_decision": "fit"},
        version_number=3,
    )

    assert result["docx_path"].endswith(f"acme_inc_ml_engineer_{str(trace_id)[:8]}_v3.docx")
    assert result["selected_resume_track"] == "resume_track_python_ml"
    assert result["source_docx_path"].endswith("resume_track_python_ml.docx")
    output_doc = Document(result["docx_path"])
    all_text = "\n".join(p.text for p in output_doc.paragraphs)
    assert "Updated summary with Python and LLM impact." in all_text
    assert "Python, SQL" in all_text


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "model_text",
    [
        json.dumps(
            {
                "edited_sections": {"summary": "Updated summary"},
                "changes_applied": ["Updated summary"],
                "evaluation": {"ats_score_before": 50, "ats_score_after": 70, "checklist_passed": True},
            }
        ),
        "```json\n"
        + json.dumps(
            {
                "edited_sections": {"summary": "Updated summary"},
                "changes_applied": ["Updated summary"],
                "evaluation": {"ats_score_before": 50, "ats_score_after": 70, "checklist_passed": True},
            }
        )
        + "\n```",
    ],
)
async def test_run_handles_plain_and_fenced_json(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    model_text: str,
) -> None:
    agent = _build_agent(monkeypatch, tmp_path)
    agent._call_model = AsyncMock(return_value={"text": model_text})  # type: ignore[method-assign]

    result = await agent.run(
        research_output=_research_output(),
        trace_id=uuid4(),
        job_context={"company": "Acme", "job_title": "ML Engineer", "relevance_decision": "fit"},
        version_number=1,
    )
    assert result["ats_score_after"] == 70


@pytest.mark.asyncio
async def test_run_invalid_schema_raises(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    agent = _build_agent(monkeypatch, tmp_path)
    agent._call_model = AsyncMock(return_value={"text": json.dumps({"edited_sections": {}})})  # type: ignore[method-assign]
    with pytest.raises(ValueError, match="missing required keys"):
        await agent.run(
            research_output=_research_output(),
            trace_id=uuid4(),
            job_context={"company": "Acme", "job_title": "ML Engineer", "relevance_decision": "fit"},
            version_number=1,
        )


@pytest.mark.asyncio
async def test_run_missing_track_docx_raises(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    agent = _build_agent(monkeypatch, tmp_path)
    track_docx_path = Path(agent.settings.resume_docx_tracks_dir) / "resume_track_python_ml.docx"
    track_docx_path.unlink(missing_ok=True)
    agent._call_model = AsyncMock(return_value={"text": "{}"})  # type: ignore[method-assign]
    with pytest.raises(FileNotFoundError, match="Editable source DOCX not found"):
        await agent.run(
            research_output=_research_output(),
            trace_id=uuid4(),
            job_context={"company": "Acme", "job_title": "ML Engineer", "relevance_decision": "fit"},
            version_number=1,
        )


@pytest.mark.asyncio
async def test_run_uses_ats_scorer_fallback_when_after_missing(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    agent = _build_agent(monkeypatch, tmp_path)
    model_payload = {
        "edited_sections": {"summary": "Updated summary"},
        "changes_applied": ["Updated summary"],
        "evaluation": {"ats_score_before": 50, "checklist_passed": True},
    }
    agent._call_model = AsyncMock(return_value={"text": json.dumps(model_payload)})  # type: ignore[method-assign]

    result = await agent.run(
        research_output=_research_output(),
        trace_id=uuid4(),
        job_context={"company": "Acme", "job_title": "ML Engineer", "relevance_decision": "fit"},
        version_number=1,
    )
    assert result["ats_score_after"] == 77


