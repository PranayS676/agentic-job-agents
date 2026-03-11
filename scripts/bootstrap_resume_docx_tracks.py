from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document

from job_agent_runtime.agents.resume_tracks import load_resume_tracks
from job_platform.config import get_settings


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Bootstrap per-track editable DOCX files from BASE_RESUME_DOCX."
    )
    parser.add_argument(
        "--resume-tracks-dir",
        type=Path,
        default=None,
        help="Directory containing normalized resume track JSON files. Defaults to RESUME_TRACKS_DIR.",
    )
    parser.add_argument(
        "--resume-docx-tracks-dir",
        type=Path,
        default=None,
        help="Directory where per-track DOCX files will be written. Defaults to RESUME_DOCX_TRACKS_DIR.",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite existing per-track DOCX files.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    settings = get_settings()

    if settings.base_resume_docx is None:
        raise SystemExit("BASE_RESUME_DOCX is not configured.")

    base_resume_docx = settings.resolve_path(settings.base_resume_docx)
    if not base_resume_docx.is_file():
        raise SystemExit(f"BASE_RESUME_DOCX not found: {base_resume_docx}")

    if args.resume_tracks_dir is not None:
        resume_tracks_dir = args.resume_tracks_dir
    elif settings.resume_tracks_dir is not None:
        resume_tracks_dir = settings.resolve_path(settings.resume_tracks_dir)
    else:
        raise SystemExit("RESUME_TRACKS_DIR is not configured and --resume-tracks-dir was not provided.")

    if args.resume_docx_tracks_dir is not None:
        resume_docx_tracks_dir = args.resume_docx_tracks_dir
    elif settings.resume_docx_tracks_dir is not None:
        resume_docx_tracks_dir = settings.resolve_path(settings.resume_docx_tracks_dir)
    else:
        raise SystemExit(
            "RESUME_DOCX_TRACKS_DIR is not configured and --resume-docx-tracks-dir was not provided."
        )

    tracks = load_resume_tracks(resume_tracks_dir)
    if not tracks:
        raise SystemExit(f"No resume track JSON files found in {resume_tracks_dir}")

    resume_docx_tracks_dir.mkdir(parents=True, exist_ok=True)
    base_docx_is_valid = _is_valid_docx(base_resume_docx)
    written_paths: list[str] = []
    skipped_paths: list[str] = []
    for track in tracks:
        target_path = resume_docx_tracks_dir / f"{track['track_id']}.docx"
        if target_path.exists() and not args.force:
            skipped_paths.append(target_path.name)
            continue
        if base_docx_is_valid:
            shutil.copyfile(base_resume_docx, target_path)
        else:
            _build_docx_from_track(track, target_path)
        written_paths.append(target_path.name)

    summary = {
        "base_resume_docx": str(base_resume_docx),
        "base_resume_docx_valid": base_docx_is_valid,
        "resume_tracks_dir": str(resume_tracks_dir),
        "resume_docx_tracks_dir": str(resume_docx_tracks_dir),
        "track_count": len(tracks),
        "written": written_paths,
        "skipped": skipped_paths,
    }
    print(json.dumps(summary, indent=2))


def _is_valid_docx(path: Path) -> bool:
    try:
        Document(str(path))
    except Exception:  # noqa: BLE001
        return False
    return True


def _build_docx_from_track(track: dict, target_path: Path) -> None:
    document = Document()
    sections = track.get("sections", {})

    document.add_heading("Summary", level=1)
    document.add_paragraph(str(sections.get("summary") or ""))

    document.add_heading("Skills", level=1)
    document.add_paragraph(str(sections.get("skills") or ""))

    experience_recent = str(sections.get("experience_recent_role") or "")
    if experience_recent:
        document.add_heading("Relevant Experience", level=1)
        for line in experience_recent.splitlines():
            cleaned = line.strip()
            if cleaned:
                document.add_paragraph(cleaned)

    for key in sorted(section for section in sections if str(section).startswith("experience_prior_role_")):
        document.add_heading("Prior Experience", level=1)
        for line in str(sections[key]).splitlines():
            cleaned = line.strip()
            if cleaned:
                document.add_paragraph(cleaned)

    education = str(sections.get("education") or "")
    if education:
        document.add_heading("Education", level=1)
        for line in education.splitlines():
            cleaned = line.strip()
            if cleaned:
                document.add_paragraph(cleaned)

    target_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target_path))


if __name__ == "__main__":
    main()
